#!/usr/bin/env python3
"""Converte manuscript.md em manuscript.docx no formato de submissão.

    pip install python-docx
    python3 docs/paper/gerar_docx.py [entrada.md] [saida.docx]

Formato aplicado: Times New Roman 12, espaço duplo, numeração de linhas e
numeração de página — o que a maioria dos periódicos da Elsevier pede para
revisão. Tabelas em Markdown viram tabelas do Word; **negrito**, *itálico* e
`código` são preservados.
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("Instale a dependência primeiro:  pip install python-docx")

RAIZ = Path(__file__).resolve().parent
ENTRADA = Path(sys.argv[1]) if len(sys.argv) > 1 else RAIZ / "manuscript.md"
SAIDA = Path(sys.argv[2]) if len(sys.argv) > 2 else RAIZ / "manuscript.docx"

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\^[^^]+\^)")


def escreve_inline(par, texto):
    """Aplica negrito, itálico, monoespaçado e sobrescrito dentro do parágrafo."""
    for pedaco in INLINE.split(texto):
        if not pedaco:
            continue
        if pedaco.startswith("**") and pedaco.endswith("**"):
            par.add_run(pedaco[2:-2]).bold = True
        elif pedaco.startswith("*") and pedaco.endswith("*") and len(pedaco) > 2:
            par.add_run(pedaco[1:-1]).italic = True
        elif pedaco.startswith("`") and pedaco.endswith("`"):
            r = par.add_run(pedaco[1:-1])
            r.font.name = "Courier New"
        elif pedaco.startswith("^") and pedaco.endswith("^"):
            r = par.add_run(pedaco[1:-1])
            r.font.superscript = True
        else:
            par.add_run(pedaco)


def numeracao_de_linha(secao):
    """Numeração contínua de linhas — exigida por muitos periódicos na revisão."""
    sectPr = secao._sectPr
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:distance"), "360")
    sectPr.append(lnNumType)


def numero_de_pagina(doc):
    rodape = doc.sections[0].footer.paragraphs[0]
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rodape.add_run()
    for instrucao, valor in (("begin", None), (None, "PAGE"), ("end", None)):
        el = OxmlElement("w:fldChar") if instrucao else OxmlElement("w:instrText")
        if instrucao:
            el.set(qn("w:fldCharType"), instrucao)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = valor
        r._r.append(el)


def main():
    linhas = ENTRADA.read_text(encoding="utf-8").split("\n")
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)
    numeracao_de_linha(doc.sections[0])
    numero_de_pagina(doc)

    i = 0
    n_tabelas = 0
    while i < len(linhas):
        linha = linhas[i].rstrip()

        if not linha.strip():
            i += 1
            continue

        # separador horizontal → quebra de página (as seções do manuscrito)
        if re.fullmatch(r"-{3,}", linha.strip()):
            doc.add_page_break()
            i += 1
            continue

        # tabela
        if linha.lstrip().startswith("|") and i + 1 < len(linhas) and re.match(r"^\s*\|[\s:|-]+\|\s*$", linhas[i + 1]):
            corpo = []
            cab = [c.strip() for c in linha.strip().strip("|").split("|")]
            i += 2
            while i < len(linhas) and linhas[i].lstrip().startswith("|"):
                corpo.append([c.strip() for c in linhas[i].strip().strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1, cols=len(cab))
            t.style = "Table Grid"
            for j, texto in enumerate(cab):
                cel = t.rows[0].cells[j]
                cel.text = ""
                p = cel.paragraphs[0]
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                escreve_inline(p, texto)
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
            for linha_dados in corpo:
                cels = t.add_row().cells
                for j, texto in enumerate(linha_dados[: len(cab)]):
                    cels[j].text = ""
                    p = cels[j].paragraphs[0]
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    escreve_inline(p, texto)
                    for r in p.runs:
                        r.font.size = Pt(10)
            n_tabelas += 1
            doc.add_paragraph()
            continue

        # cabeçalhos
        m = re.match(r"^(#{1,4})\s+(.*)$", linha)
        if m:
            nivel = len(m.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            escreve_inline(p, m.group(2))
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14 - nivel)
            i += 1
            continue

        # citação em bloco
        if linha.lstrip().startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            escreve_inline(p, linha.lstrip().lstrip(">").strip())
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # listas
        m = re.match(r"^\s*[-*]\s+(.*)$", linha)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            escreve_inline(p, m.group(1))
            i += 1
            continue
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", linha)
        if m:
            p = doc.add_paragraph(style="List Number")
            escreve_inline(p, m.group(2))
            i += 1
            continue

        # parágrafo comum, juntando linhas até a próxima em branco
        bloco = [linha]
        i += 1
        while i < len(linhas) and linhas[i].strip() and not re.match(r"^(#|\||>|\s*[-*]\s|\s*\d+\.\s|-{3,})", linhas[i]):
            bloco.append(linhas[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escreve_inline(p, " ".join(bloco))

    doc.save(SAIDA)
    n_palavras = len(ENTRADA.read_text(encoding="utf-8").split())
    print(f"{SAIDA}  —  {n_tabelas} tabela(s), ~{n_palavras} palavras no fonte")


if __name__ == "__main__":
    main()
